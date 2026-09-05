import io

import pytest
from docx import Document as DocxDocument

from app.connectors.document import extraction
from app.core.exceptions import ConnectorConfigError


def test_extract_text_decodes_plain_text_files():
    result = extraction.extract_text(filename="note.txt", raw_bytes=b"Bonjour le monde")

    assert result == "Bonjour le monde"


def test_extract_text_rejects_oversized_files(monkeypatch: pytest.MonkeyPatch):
    monkeypatch.setattr(extraction, "MAX_FILE_SIZE_BYTES", 10)

    with pytest.raises(ConnectorConfigError):
        extraction.extract_text(filename="note.txt", raw_bytes=b"0123456789ABCDEF")


def test_extract_text_rejects_unsupported_binary_format():
    with pytest.raises(ConnectorConfigError):
        extraction.extract_text(filename="image.png", raw_bytes=b"\x89PNG\r\n\x1a\n\x00\x01")


def test_extract_text_reads_docx_paragraphs():
    document = DocxDocument()
    document.add_paragraph("Premier paragraphe.")
    document.add_paragraph("Second paragraphe.")
    buffer = io.BytesIO()
    document.save(buffer)

    result = extraction.extract_text(filename="note.docx", raw_bytes=buffer.getvalue())

    assert "Premier paragraphe." in result
    assert "Second paragraphe." in result


def test_extract_text_reads_pdf_pages(monkeypatch: pytest.MonkeyPatch):
    class FakePage:
        def __init__(self, text: str):
            self._text = text

        def extract_text(self) -> str:
            return self._text

    class FakeReader:
        def __init__(self, _stream: io.BytesIO):
            self.pages = [FakePage("Page un."), FakePage("Page deux.")]

    monkeypatch.setattr(extraction, "PdfReader", FakeReader)

    result = extraction.extract_text(filename="note.pdf", raw_bytes=b"%PDF-1.4 fake")

    assert result == "Page un.\n\nPage deux."


def test_extract_text_raises_when_pdf_has_no_extractable_text(monkeypatch: pytest.MonkeyPatch):
    class FakePage:
        def extract_text(self) -> str:
            return ""

    class FakeReader:
        def __init__(self, _stream: io.BytesIO):
            self.pages = [FakePage()]

    monkeypatch.setattr(extraction, "PdfReader", FakeReader)

    with pytest.raises(ConnectorConfigError):
        extraction.extract_text(filename="scan.pdf", raw_bytes=b"%PDF-1.4 fake")
